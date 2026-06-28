"""
Real Madrid CF — Analysis Report Generator
Produces PDF and DOCX reports strictly from dashboard metrics and user selections.
No synthetic data. Every value is traceable to the underlying dataset.
"""

from __future__ import annotations
import io
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any

import pandas as pd

# ── ReportLab (PDF) ───────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm, mm
from reportlab.lib import colors as rl_colors
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.platypus import (
    BaseDocTemplate, Frame, PageTemplate, Paragraph, Spacer,
    Table, TableStyle, Image as RLImage, PageBreak, HRFlowable,
    KeepTogether, NextPageTemplate,
)

# ── python-docx (DOCX) ────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Optional Plotly image export ──────────────────────────────────────────────
try:
    import plotly.io as _pio
    _PLOTLY_OK = True
except ImportError:
    _PLOTLY_OK = False

# ── Brand ─────────────────────────────────────────────────────────────────────
_NAVY_RGB  = (11, 23, 48)
_GOLD_RGB  = (212, 175, 55)
_WHITE_RGB = (255, 255, 255)
_GRAY_RGB  = (100, 116, 139)
_GREEN_RGB = (5, 150, 105)
_RED_RGB   = (220, 38, 38)
_AMBER_RGB = (217, 119, 6)
_BLUE_RGB  = (37, 99, 235)

_NAVY_RL  = rl_colors.Color(*[v/255 for v in _NAVY_RGB])
_GOLD_RL  = rl_colors.Color(*[v/255 for v in _GOLD_RGB])
_WHITE_RL = rl_colors.white
_GRAY_RL  = rl_colors.Color(*[v/255 for v in _GRAY_RGB])
_GREEN_RL = rl_colors.Color(*[v/255 for v in _GREEN_RGB])
_RED_RL   = rl_colors.Color(*[v/255 for v in _RED_RGB])
_AMBER_RL = rl_colors.Color(*[v/255 for v in _AMBER_RGB])
_BLUE_RL  = rl_colors.Color(*[v/255 for v in _BLUE_RGB])

_CREST_PATH   = str(Path(__file__).parent / "dashboard" / "app" / "assets" / "rm_crest.jpg")
_UE_LOGO_PATH = str(Path(__file__).parent / "dashboard" / "app" / "assets" / "ue_real_madrid_logo.png")
PAGE_W, PAGE_H = A4


# ═════════════════════════════════════════════════════════════════════════════
# HELPERS: figure → PNG bytes
# ═════════════════════════════════════════════════════════════════════════════

def _fig_to_png(fig, width: int = 720, height: int = 360) -> bytes | None:
    """Convert a Plotly figure to PNG bytes, returns None on failure."""
    if not _PLOTLY_OK or fig is None:
        return None
    try:
        return _pio.to_image(fig, format="png", width=width, height=height, scale=1.5)
    except Exception:
        return None


# ═════════════════════════════════════════════════════════════════════════════
# REPORT CONFIG DATACLASS (plain dict, keeps it simple)
# ═════════════════════════════════════════════════════════════════════════════

def build_report_config(
    competition: str,
    season: str,
    match_mode: str,
    fps: list[str],
    match_labels: list[str],
    players: list[str],
    match_records: list[dict],
    squad_df: pd.DataFrame,
    player_dfs: dict[str, pd.DataFrame],
    figures: dict[str, Any],
) -> dict:
    """
    Assembles a self-contained config dict from real dashboard data.
    All numeric values come from the caller (dashboard functions).
    """
    return {
        "competition":    competition,
        "season":         season,
        "match_mode":     match_mode,
        "fps":            fps,
        "match_labels":   match_labels,
        "players":        players,
        "match_records":  match_records,   # list of {meta, kpis, ppda, player_stats}
        "squad_df":       squad_df,
        "player_dfs":     player_dfs,      # {name: DataFrame}
        "figures":        figures,          # {key: Plotly figure}
        "report_id":      uuid.uuid4().hex[:8].upper(),
        "generated_at":   datetime.now(),
        "crest_path":     _CREST_PATH,
    }


# ═════════════════════════════════════════════════════════════════════════════
# DATA SUMMARY HELPERS (derive all text from real numbers)
# ═════════════════════════════════════════════════════════════════════════════

def _exec_summary_lines(cfg: dict) -> list[str]:
    """Generate executive summary bullet lines strictly from calculated metrics."""
    recs = cfg["match_records"]
    if not recs:
        return ["No match data available for the selected filters."]

    n  = len(recs)
    w  = sum(1 for r in recs if r["meta"].get("rm_score", 0) > r["meta"].get("opp_score", 0))
    d  = sum(1 for r in recs if r["meta"].get("rm_score", 0) == r["meta"].get("opp_score", 0))
    l  = n - w - d
    gf = sum(r["meta"].get("rm_score",  0) for r in recs)
    ga = sum(r["meta"].get("opp_score", 0) for r in recs)

    xg_for  = sum(r["kpis"].get("xg_for",     0) for r in recs)
    xg_ag   = sum(r["kpis"].get("xg_against", 0) for r in recs)
    poss    = [r["kpis"].get("possession",    0) for r in recs if r["kpis"].get("possession")]
    shots   = sum(r["kpis"].get("shots_total", 0) for r in recs)
    pass_ac = [r["kpis"].get("pass_accuracy", 0) for r in recs if r["kpis"].get("pass_accuracy")]
    ppda_v  = [r["ppda"] for r in recs if isinstance(r.get("ppda"), (int, float))]

    lines = [
        f"Analysis covers {n} match{'es' if n > 1 else ''} in {cfg['competition']} {cfg['season']}.",
        f"Record: {w}W / {d}D / {l}L — {gf} goals scored, {ga} conceded.",
    ]
    if n > 0:
        lines.append(
            f"Total xG generated: {xg_for:.2f} (avg {xg_for/n:.2f}/match) vs "
            f"{xg_ag:.2f} xG conceded (avg {xg_ag/n:.2f}/match)."
        )
    if shots:
        lines.append(f"Total shots on goal: {shots} across {n} match{'es' if n > 1 else ''}.")
    if poss:
        lines.append(f"Average possession: {sum(poss)/len(poss):.1f}%.")
    if pass_ac:
        lines.append(f"Average pass accuracy: {sum(pass_ac)/len(pass_ac):.1f}%.")
    if ppda_v:
        avg_ppda = sum(ppda_v) / len(ppda_v)
        press_desc = "high" if avg_ppda < 8 else ("moderate" if avg_ppda < 12 else "low")
        lines.append(f"Average PPDA: {avg_ppda:.2f} — indicating {press_desc} pressing intensity.")

    if cfg["players"]:
        lines.append(
            f"Player analysis includes: {', '.join(cfg['players'])}."
        )
    return lines


def _pre_match_lines(cfg: dict) -> list[str]:
    """Pre-match: form, benchmark indicators from available data before first selected match."""
    recs  = cfg["match_records"]
    squad = cfg["squad_df"]
    lines = []

    if squad is not None and not squad.empty:
        top_scorers = (squad.sort_values("goals", ascending=False)
                       .head(3)[["player_name", "goals", "matches_played"]]
                       .to_dict("records") if "goals" in squad.columns else [])
        if top_scorers:
            ts_str = "; ".join(
                f"{r.get('player_name','?')} ({r.get('goals',0)}G in {r.get('matches_played',0)} apps)"
                for r in top_scorers
            )
            lines.append(f"Top scorers in squad pool: {ts_str}.")

        if "xg" in squad.columns and "matches_played" in squad.columns:
            total_squad_xg = squad["xg"].sum()
            lines.append(f"Squad cumulative xG: {total_squad_xg:.2f} across all tracked matches.")

    if recs:
        first = recs[0]
        meta  = first["meta"]
        lines.append(
            f"Selected match period opens with: {meta.get('home_team', '?')} vs "
            f"{meta.get('away_team', '?')} ({meta.get('date', '?')}, {meta.get('competition', '?')})."
        )
    if not lines:
        lines = ["Insufficient historical data for pre-match benchmark generation."]
    return lines


def _match_analysis_lines(cfg: dict) -> list[str]:
    """Match analysis: what happened — purely from event/KPI data."""
    recs  = cfg["match_records"]
    lines = []
    for r in recs:
        meta = r["meta"]
        kpis = r["kpis"]
        score = meta.get("score_str", "?-?")
        result_word = (
            "won" if meta.get("rm_score", 0) > meta.get("opp_score", 0) else
            "drew" if meta.get("rm_score", 0) == meta.get("opp_score", 0) else "lost"
        )
        lines.append(
            f"{meta.get('home_team','RM')} vs {meta.get('away_team','Opp')} "
            f"({meta.get('date','?')}): Real Madrid {result_word} {score}."
        )
        if kpis.get("possession"):
            lines.append(
                f"  Possession: {kpis['possession']}% | "
                f"Pass accuracy: {kpis.get('pass_accuracy', '?')}% | "
                f"Shots: {kpis.get('shots_total', '?')}."
            )
        if kpis.get("xg_for") is not None:
            lines.append(
                f"  xG For: {kpis['xg_for']:.2f} | xG Against: {kpis.get('xg_against', 0):.2f} | "
                f"PPDA: {r.get('ppda', '?')}."
            )
    if not lines:
        lines = ["No match event data available for the selected period."]
    return lines


def _post_match_lines(cfg: dict) -> list[str]:
    """Post-match: outcome evaluation from KPIs vs what happened."""
    recs  = cfg["match_records"]
    if not recs:
        return ["No data available for post-match analysis."]

    n   = len(recs)
    gf  = sum(r["meta"].get("rm_score", 0) for r in recs)
    ga  = sum(r["meta"].get("opp_score", 0) for r in recs)
    xgf = sum(r["kpis"].get("xg_for", 0) for r in recs)
    xga = sum(r["kpis"].get("xg_against", 0) for r in recs)
    lines = [
        f"Goals scored ({gf}) vs xG generated ({xgf:.2f}): "
        + ("outperformed xG model." if gf > xgf else
           "underperformed xG model." if gf < xgf else "matched xG model exactly."),
        f"Goals conceded ({ga}) vs xG against ({xga:.2f}): "
        + ("defensive overperformance." if ga < xga else
           "defensive underperformance." if ga > xga else "matched expected goals against."),
    ]
    poss_vals = [r["kpis"].get("possession", 0) for r in recs if r["kpis"].get("possession")]
    if poss_vals:
        avg_p = sum(poss_vals) / len(poss_vals)
        lines.append(
            f"Average possession {avg_p:.1f}% — "
            + ("dominant ball control." if avg_p > 55 else
               "contested possession battle." if avg_p >= 45 else "lower possession profile.")
        )
    ppda_v = [r["ppda"] for r in recs if isinstance(r.get("ppda"), (int, float))]
    if ppda_v:
        avg_ppda = sum(ppda_v) / len(ppda_v)
        lines.append(
            f"PPDA average {avg_ppda:.2f} — "
            + ("high pressing output, consistent with proactive pressing scheme."
               if avg_ppda < 8 else
               "moderate press; scope to intensify in future matches."
               if avg_ppda < 12 else "lower pressing intensity observed.")
        )
    return lines


def _player_summary_lines(cfg: dict) -> list[str]:
    """Per-player summary from per-match DataFrames."""
    lines = []
    for pname, pdf in cfg["player_dfs"].items():
        if pdf is None or pdf.empty:
            lines.append(f"{pname}: No per-match data available.")
            continue
        nm    = len(pdf)
        goals = pdf.get("goals", pd.Series(dtype=float)).sum() if "goals" in pdf.columns else 0
        shots = pdf.get("shots", pd.Series(dtype=float)).sum() if "shots" in pdf.columns else 0
        passes= pdf.get("passes",pd.Series(dtype=float)).sum() if "passes" in pdf.columns else 0
        xg    = pdf.get("xg",   pd.Series(dtype=float)).sum() if "xg" in pdf.columns else 0
        lines.append(
            f"{pname} ({nm} match{'es' if nm > 1 else ''}): "
            f"{int(goals)} goals, {int(shots)} shots, {int(passes)} passes, "
            f"{float(xg):.2f} xG."
        )
    return lines if lines else ["No player data for selection."]


# ═════════════════════════════════════════════════════════════════════════════
# PDF GENERATION
# ═════════════════════════════════════════════════════════════════════════════

def _pdf_styles() -> dict:
    """Return a dict of named ParagraphStyle objects."""
    base = dict(fontName="Helvetica", leading=14)
    return {
        "cover_title": ParagraphStyle(
            "cover_title", fontName="Helvetica-Bold",
            fontSize=28, textColor=_GOLD_RL, alignment=TA_CENTER, leading=34,
        ),
        "cover_sub": ParagraphStyle(
            "cover_sub", fontName="Helvetica",
            fontSize=13, textColor=_WHITE_RL, alignment=TA_CENTER, leading=18,
        ),
        "cover_meta": ParagraphStyle(
            "cover_meta", fontName="Helvetica",
            fontSize=10, textColor=_GRAY_RL, alignment=TA_CENTER, leading=14,
        ),
        "h1": ParagraphStyle(
            "h1", fontName="Helvetica-Bold",
            fontSize=16, textColor=_NAVY_RL, spaceBefore=14, spaceAfter=6,
        ),
        "h2": ParagraphStyle(
            "h2", fontName="Helvetica-Bold",
            fontSize=12, textColor=_GOLD_RL, spaceBefore=10, spaceAfter=4,
        ),
        "h3": ParagraphStyle(
            "h3", fontName="Helvetica-Bold",
            fontSize=10, textColor=_NAVY_RL, spaceBefore=8, spaceAfter=2,
        ),
        "body": ParagraphStyle(
            "body", fontName="Helvetica",
            fontSize=9, textColor=rl_colors.Color(0.12, 0.18, 0.27),
            leading=13, spaceAfter=4,
        ),
        "bullet": ParagraphStyle(
            "bullet", fontName="Helvetica",
            fontSize=9, textColor=rl_colors.Color(0.12, 0.18, 0.27),
            leading=13, leftIndent=12, bulletIndent=0, spaceAfter=3,
        ),
        "caption": ParagraphStyle(
            "caption", fontName="Helvetica-Oblique",
            fontSize=8, textColor=_GRAY_RL, alignment=TA_CENTER, spaceAfter=4,
        ),
        "label": ParagraphStyle(
            "label", fontName="Helvetica-Bold",
            fontSize=8, textColor=_NAVY_RL,
        ),
    }


def _pdf_table_style(header_bg=None) -> TableStyle:
    bg = header_bg or _NAVY_RL
    return TableStyle([
        ("BACKGROUND",   (0, 0), (-1, 0),  bg),
        ("TEXTCOLOR",    (0, 0), (-1, 0),  _WHITE_RL),
        ("FONTNAME",     (0, 0), (-1, 0),  "Helvetica-Bold"),
        ("FONTSIZE",     (0, 0), (-1, 0),  8),
        ("ALIGN",        (0, 0), (-1, -1), "CENTER"),
        ("VALIGN",       (0, 0), (-1, -1), "MIDDLE"),
        ("FONTNAME",     (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE",     (0, 1), (-1, -1), 8),
        ("ROWBACKGROUNDS",(0, 1),(-1, -1), [rl_colors.white,
                                            rl_colors.Color(0.96, 0.97, 0.99)]),
        ("GRID",         (0, 0), (-1, -1), 0.4, _GRAY_RL),
        ("TOPPADDING",   (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 4),
        ("LEFTPADDING",  (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
    ])


def _df_to_pdf_table(df: pd.DataFrame, col_widths=None) -> Table | None:
    if df is None or df.empty:
        return None
    data = [list(df.columns)] + [[str(v) for v in row] for _, row in df.iterrows()]
    t = Table(data, colWidths=col_widths, repeatRows=1)
    t.setStyle(_pdf_table_style())
    return t


def _divider(color=None) -> HRFlowable:
    return HRFlowable(width="100%", thickness=1, color=color or _GOLD_RL, spaceAfter=6)


class _NavyPageTemplate(PageTemplate):
    """Page template with navy header bar and gold footer line."""

    def __init__(self, doc, show_cover=False):
        self.show_cover = show_cover
        frame = Frame(
            doc.leftMargin, doc.bottomMargin,
            doc.width, doc.height,
            id="normal",
        )
        super().__init__(id="main", frames=[frame])

    def beforeDrawPage(self, canvas, doc):
        w, h = PAGE_W, PAGE_H
        canvas.saveState()
        if self.show_cover:
            canvas.setFillColorRGB(*[v/255 for v in _NAVY_RGB])
            canvas.rect(0, 0, w, h, fill=1, stroke=0)
        else:
            # Header bar
            canvas.setFillColorRGB(*[v/255 for v in _NAVY_RGB])
            canvas.rect(0, h - 22*mm, w, 22*mm, fill=1, stroke=0)
            # Gold accent line below header
            canvas.setFillColorRGB(*[v/255 for v in _GOLD_RGB])
            canvas.rect(0, h - 23*mm, w, 1*mm, fill=1, stroke=0)
            # UE Real Madrid logo — right side of header
            _ue_logo_h = 14 * mm
            _ue_logo_w = _ue_logo_h * 4.2   # approximate aspect ratio of the banner
            _ue_x = w - 1.5 * cm - _ue_logo_w
            _ue_y = h - 22 * mm + (22 * mm - _ue_logo_h) / 2
            try:
                canvas.drawImage(
                    _UE_LOGO_PATH,
                    _ue_x, _ue_y,
                    width=_ue_logo_w, height=_ue_logo_h,
                    preserveAspectRatio=True, anchor="sw",
                    mask="auto",
                )
            except Exception:
                pass
            # Header text
            canvas.setFillColorRGB(1, 1, 1)
            canvas.setFont("Helvetica-Bold", 10)
            canvas.drawString(1.5*cm, h - 15*mm, "Real Madrid CF — Tactical & Player Analytics")
            canvas.setFont("Helvetica", 8)
            canvas.setFillColorRGB(*[v/255 for v in _GOLD_RGB])
            canvas.drawString(1.5*cm, h - 20*mm,
                              f"{doc._cfg['competition']} · {doc._cfg['season']}")
            # Footer gold line
            canvas.setFillColorRGB(*[v/255 for v in _GOLD_RGB])
            canvas.rect(0, 12*mm, w, 0.5*mm, fill=1, stroke=0)
            # Footer text
            canvas.setFillColorRGB(*[v/255 for v in _GRAY_RGB])
            canvas.setFont("Helvetica", 7)
            canvas.drawString(1.5*cm, 7*mm,
                               f"Report ID: {doc._cfg['report_id']}  ·  "
                               f"Generated: {doc._cfg['generated_at'].strftime('%Y-%m-%d %H:%M')}")
            canvas.drawRightString(w - 1.5*cm, 7*mm, f"Page {doc.page}")
        canvas.restoreState()


class _RMDocTemplate(BaseDocTemplate):
    def __init__(self, buf, cfg):
        self._cfg = cfg
        super().__init__(
            buf,
            pagesize=A4,
            leftMargin=1.8*cm, rightMargin=1.8*cm,
            topMargin=2.8*cm,  bottomMargin=2.2*cm,
        )
        self.addPageTemplates([
            _NavyPageTemplate(self, show_cover=False),
        ])


def generate_pdf(cfg: dict) -> bytes:
    """Build and return a complete PDF report as bytes."""
    buf = io.BytesIO()
    doc = _RMDocTemplate(buf, cfg)
    S   = _pdf_styles()
    W   = doc.width
    story = []

    # ── COVER PAGE ───────────────────────────────────────────────────────────
    # Full-navy cover: draw directly, then add content as flowables with custom bg
    story.append(Spacer(1, 2*cm))
    # UE Real Madrid logo on cover
    try:
        ue_logo = RLImage(_UE_LOGO_PATH, width=8*cm, height=2.2*cm)
        ue_logo.hAlign = "CENTER"
        story.append(ue_logo)
    except Exception:
        pass
    story.append(Spacer(1, 0.8*cm))
    # RM Crest image
    try:
        crest = RLImage(cfg["crest_path"], width=3.5*cm, height=3.5*cm)
        crest.hAlign = "CENTER"
        story.append(crest)
    except Exception:
        pass
    story.append(Spacer(1, 0.6*cm))
    story.append(Paragraph("REAL MADRID CF", ParagraphStyle(
        "ct2", fontName="Helvetica-Bold", fontSize=22,
        textColor=_GOLD_RL, alignment=TA_CENTER)))
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph("Tactical &amp; Player Analytics Report", ParagraphStyle(
        "ct3", fontName="Helvetica", fontSize=14,
        textColor=_WHITE_RL, alignment=TA_CENTER)))
    story.append(Spacer(1, 1.2*cm))
    story.append(_divider(_GOLD_RL))
    story.append(Spacer(1, 0.4*cm))

    cover_meta = [
        ("Competition", cfg["competition"]),
        ("Season",      cfg["season"]),
        ("Match Selection", cfg["match_mode"]),
        ("Matches",     ", ".join(cfg["match_labels"][:3])
                        + (f" (+{len(cfg['match_labels'])-3} more)"
                           if len(cfg["match_labels"]) > 3 else "")),
        ("Athletes",    ", ".join(cfg["players"]) if cfg["players"] else "All Squad"),
        ("Report ID",   cfg["report_id"]),
        ("Generated",   cfg["generated_at"].strftime("%d %B %Y  %H:%M")),
    ]
    cover_table = Table(
        [[k + ":", v] for k, v in cover_meta],
        colWidths=[5*cm, W - 5*cm],
    )
    cover_table.setStyle(TableStyle([
        ("FONTNAME",  (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTNAME",  (1, 0), (1, -1), "Helvetica"),
        ("FONTSIZE",  (0, 0), (-1, -1), 10),
        ("TEXTCOLOR", (0, 0), (0, -1), _GOLD_RL),
        ("TEXTCOLOR", (1, 0), (1, -1), _WHITE_RL),
        ("ALIGN",     (0, 0), (0, -1), "RIGHT"),
        ("ALIGN",     (1, 0), (1, -1), "LEFT"),
        ("TOPPADDING",(0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(cover_table)
    story.append(PageBreak())

    # ── SECTION 1: FILTER SUMMARY ────────────────────────────────────────────
    story.append(Paragraph("Section 1 — Filter Summary", S["h1"]))
    story.append(_divider())
    filters = [
        ["Filter", "Value"],
        ["Competition",  cfg["competition"]],
        ["Season",       cfg["season"]],
        ["Match Mode",   cfg["match_mode"]],
        ["Matches",      "\n".join(cfg["match_labels"]) or "All"],
        ["Athletes",     "\n".join(cfg["players"]) or "All Squad"],
        ["Report ID",    cfg["report_id"]],
        ["Generated At", cfg["generated_at"].strftime("%Y-%m-%d %H:%M:%S")],
    ]
    t = Table(filters, colWidths=[5*cm, W - 5*cm])
    t.setStyle(_pdf_table_style())
    story.append(t)
    story.append(Spacer(1, 0.4*cm))

    # ── SECTION 2: EXECUTIVE SUMMARY ─────────────────────────────────────────
    story.append(Paragraph("Section 2 — Executive Summary", S["h1"]))
    story.append(_divider())
    for line in _exec_summary_lines(cfg):
        story.append(Paragraph(f"• {line}", S["bullet"]))
    story.append(Spacer(1, 0.3*cm))

    # ── SECTION 3: TEAM PERFORMANCE ──────────────────────────────────────────
    story.append(Paragraph("Section 3 — Team Performance Analysis", S["h1"]))
    story.append(_divider())
    recs = cfg["match_records"]
    if recs:
        n   = len(recs)
        gf  = sum(r["meta"].get("rm_score", 0) for r in recs)
        ga  = sum(r["meta"].get("opp_score", 0) for r in recs)
        xgf = sum(r["kpis"].get("xg_for",       0) for r in recs)
        xga = sum(r["kpis"].get("xg_against",   0) for r in recs)
        sht = sum(r["kpis"].get("shots_total",   0) for r in recs)
        poss_v = [r["kpis"].get("possession",   0) for r in recs if r["kpis"].get("possession")]
        pacc_v = [r["kpis"].get("pass_accuracy",0) for r in recs if r["kpis"].get("pass_accuracy")]
        ppda_v = [r["ppda"] for r in recs if isinstance(r.get("ppda"), (int, float))]

        team_data = [
            ["Metric",             "Total / Avg",      "Per Match"],
            ["Matches Analysed",   str(n),              "—"],
            ["Goals Scored",       str(gf),             f"{gf/n:.2f}"],
            ["Goals Conceded",     str(ga),             f"{ga/n:.2f}"],
            ["xG Generated",       f"{xgf:.2f}",        f"{xgf/n:.2f}"],
            ["xG Conceded",        f"{xga:.2f}",        f"{xga/n:.2f}"],
            ["xG Difference",      f"{(xgf-xga):+.2f}", f"{(xgf-xga)/n:+.2f}"],
            ["Shots",              str(sht),             f"{sht/n:.1f}"],
            ["Avg Possession",     f"{sum(poss_v)/len(poss_v):.1f}%" if poss_v else "N/A", "—"],
            ["Avg Pass Accuracy",  f"{sum(pacc_v)/len(pacc_v):.1f}%" if pacc_v else "N/A", "—"],
            ["Avg PPDA",           f"{sum(ppda_v)/len(ppda_v):.2f}" if ppda_v else "N/A", "—"],
        ]
        t = Table(team_data, colWidths=[7*cm, 4*cm, 4*cm])
        t.setStyle(_pdf_table_style())
        story.append(t)
    else:
        story.append(Paragraph("No team performance data for selected filters.", S["body"]))
    story.append(Spacer(1, 0.3*cm))

    # ── Per-match results table ───────────────────────────────────────────────
    if recs:
        story.append(Paragraph("Match-by-Match Results", S["h2"]))
        match_rows = [["Match", "Score", "Result", "Poss%", "PA%", "xGF", "xGA", "PPDA"]]
        for r in recs:
            meta, kpis = r["meta"], r["kpis"]
            res = ("W" if meta.get("rm_score",0) > meta.get("opp_score",0) else
                   "D" if meta.get("rm_score",0) == meta.get("opp_score",0) else "L")
            match_rows.append([
                f"{meta.get('home_team','?')} vs {meta.get('away_team','?')}",
                meta.get("score_str", "?"),
                res,
                f"{kpis.get('possession',0):.0f}%",
                f"{kpis.get('pass_accuracy',0):.0f}%",
                f"{kpis.get('xg_for',0):.2f}",
                f"{kpis.get('xg_against',0):.2f}",
                str(r.get("ppda", "?")),
            ])
        cw = [5.2*cm, 2*cm, 1.5*cm, 1.8*cm, 1.8*cm, 1.5*cm, 1.5*cm, 1.8*cm]
        mt = Table(match_rows, colWidths=cw, repeatRows=1)
        mt.setStyle(_pdf_table_style())
        story.append(mt)
    story.append(Spacer(1, 0.3*cm))

    # ── SECTION 4: PLAYER PERFORMANCE ────────────────────────────────────────
    story.append(Paragraph("Section 4 — Player Performance Analysis", S["h1"]))
    story.append(_divider())
    if cfg["players"] and cfg["player_dfs"]:
        for pname in cfg["players"]:
            pdf = cfg["player_dfs"].get(pname)
            story.append(Paragraph(pname, S["h2"]))
            if pdf is None or pdf.empty:
                story.append(Paragraph("No per-match data available.", S["body"]))
                continue
            nm     = len(pdf)
            agg    = pdf[["passes","shots","goals","tackles","xg"]].sum()
            pg     = (agg / max(nm, 1)).round(3)
            p_data = [
                ["Metric",      "Total",        "Per Match"],
                ["Matches",     str(nm),         "—"],
                ["Passes",      str(int(agg.get("passes",0))),   f"{pg.get('passes',0):.1f}"],
                ["Shots",       str(int(agg.get("shots",0))),    f"{pg.get('shots',0):.2f}"],
                ["Goals",       str(int(agg.get("goals",0))),    f"{pg.get('goals',0):.2f}"],
                ["Tackles",     str(int(agg.get("tackles",0))),  f"{pg.get('tackles',0):.2f}"],
                ["xG",          f"{float(agg.get('xg',0)):.2f}", f"{float(pg.get('xg',0)):.3f}"],
            ]
            pt = Table(p_data, colWidths=[6*cm, 4*cm, 4*cm])
            pt.setStyle(_pdf_table_style())
            story.append(pt)
            # Per-match breakdown
            if "match" in pdf.columns:
                pm_cols = [c for c in ["match","result","passes","shots","goals","xg"] if c in pdf.columns]
                if pm_cols:
                    story.append(Spacer(1, 0.2*cm))
                    story.append(Paragraph("Per-Match Breakdown", S["h3"]))
                    pm_data = [pm_cols] + [[str(v) for v in row[pm_cols].tolist()]
                                           for _, row in pdf.iterrows()]
                    pm_t = Table(pm_data, repeatRows=1)
                    pm_t.setStyle(_pdf_table_style())
                    story.append(pm_t)
            # Embed radar chart if available
            radar_fig = cfg["figures"].get(f"radar_{pname}")
            if radar_fig is None and len(cfg["players"]) >= 1:
                radar_fig = cfg["figures"].get("radar_comparison")
            png = _fig_to_png(radar_fig, width=560, height=320)
            if png:
                story.append(Spacer(1, 0.2*cm))
                img = RLImage(io.BytesIO(png), width=14*cm, height=8*cm)
                img.hAlign = "CENTER"
                story.append(img)
                story.append(Paragraph("Radar chart: normalised to % of squad maximum per metric.",
                                       S["caption"]))
            story.append(Spacer(1, 0.3*cm))
    elif not cfg["squad_df"].empty:
        # Show squad-level table when no specific players chosen
        story.append(Paragraph("Squad-Level Performance Summary", S["h2"]))
        sq = cfg["squad_df"]
        show_cols = [c for c in ["player_name","matches_played","goals","assists",
                                  "passes","shots","tackles","xg"] if c in sq.columns]
        sq_top = sq[show_cols].sort_values("goals", ascending=False).head(15)
        sq_top.columns = [c.replace("_", " ").title() for c in sq_top.columns]
        sqt = _df_to_pdf_table(sq_top)
        if sqt:
            story.append(sqt)
    else:
        story.append(Paragraph("No player data available for selected filters.", S["body"]))
    story.append(Spacer(1, 0.3*cm))

    # ── SECTION 5: TACTICAL ANALYSIS ─────────────────────────────────────────
    story.append(Paragraph("Section 5 — Tactical Analysis", S["h1"]))
    story.append(_divider())
    tac_figs = {k: v for k, v in cfg["figures"].items()
                if k in ("press_map", "press_class", "position_map", "tactical_bars")}
    if tac_figs:
        for fig_key, fig in tac_figs.items():
            label = fig_key.replace("_", " ").title()
            png = _fig_to_png(fig, width=680, height=360)
            if png:
                story.append(Paragraph(label, S["h2"]))
                img = RLImage(io.BytesIO(png), width=15*cm, height=8*cm)
                img.hAlign = "CENTER"
                story.append(img)
                story.append(Spacer(1, 0.3*cm))
    else:
        story.append(Paragraph(
            "Tactical visualisations are available in the Tactical Phases section "
            "of the dashboard. Include tactical figures by selecting them during report generation.",
            S["body"]
        ))
    story.append(Spacer(1, 0.3*cm))

    # ── SECTION 6: MATCH ANALYSIS ─────────────────────────────────────────────
    story.append(Paragraph("Section 6 — Match Analysis", S["h1"]))
    story.append(_divider())
    for line in _match_analysis_lines(cfg):
        story.append(Paragraph(line, S["body"] if not line.startswith("  ") else S["bullet"]))

    match_figs = {k: v for k, v in cfg["figures"].items()
                  if k in ("shot_map", "xg_chart", "pass_map", "pass_network")}
    for fig_key, fig in match_figs.items():
        label = fig_key.replace("_", " ").title()
        png = _fig_to_png(fig, width=680, height=360)
        if png:
            story.append(Spacer(1, 0.2*cm))
            story.append(Paragraph(label, S["h2"]))
            img = RLImage(io.BytesIO(png), width=15*cm, height=8*cm)
            img.hAlign = "CENTER"
            story.append(img)
    story.append(Spacer(1, 0.3*cm))

    # ── SECTION 7: BENCHMARKING ───────────────────────────────────────────────
    story.append(Paragraph("Section 7 — Benchmarking Analysis", S["h1"]))
    story.append(_divider())
    bm_fig = cfg["figures"].get("benchmarking_goals") or cfg["figures"].get("benchmarking_xg")
    if bm_fig:
        png = _fig_to_png(bm_fig, width=680, height=340)
        if png:
            img = RLImage(io.BytesIO(png), width=15*cm, height=7.5*cm)
            img.hAlign = "CENTER"
            story.append(img)
    else:
        story.append(Paragraph(
            "Season-level benchmarking charts are generated from aggregated competition data. "
            "Full benchmarking visuals are available in the Benchmarking section of the dashboard.",
            S["body"]
        ))
    story.append(Spacer(1, 0.3*cm))

    # ── SECTION 8: VISUALISATION APPENDIX ────────────────────────────────────
    story.append(PageBreak())
    story.append(Paragraph("Section 8 — Visualisation Appendix", S["h1"]))
    story.append(_divider())
    appendix_keys = [k for k in cfg["figures"] if k not in
                     {"shot_map","xg_chart","pass_map","pass_network",
                      "press_map","press_class","position_map","tactical_bars",
                      "benchmarking_goals","benchmarking_xg"}]
    included = 0
    for fig_key in appendix_keys:
        fig = cfg["figures"][fig_key]
        if fig is None:
            continue
        png = _fig_to_png(fig, width=640, height=340)
        if png:
            label = fig_key.replace("_", " ").title()
            story.append(Paragraph(label, S["h2"]))
            img = RLImage(io.BytesIO(png), width=15*cm, height=8*cm)
            img.hAlign = "CENTER"
            story.append(img)
            story.append(Spacer(1, 0.25*cm))
            included += 1
    if included == 0:
        story.append(Paragraph("No additional visualisations captured for this report.", S["body"]))
    story.append(Spacer(1, 0.3*cm))

    # ── CONCLUSION ───────────────────────────────────────────────────────────
    story.append(Paragraph("Conclusion", S["h1"]))
    story.append(_divider())

    story.append(Paragraph("A. Pre-Match Analysis", S["h2"]))
    for line in _pre_match_lines(cfg):
        story.append(Paragraph(f"• {line}", S["bullet"]))
    story.append(Spacer(1, 0.2*cm))

    story.append(Paragraph("B. Match Analysis", S["h2"]))
    for line in _match_analysis_lines(cfg):
        story.append(Paragraph(line, S["body"] if not line.startswith("  ") else S["bullet"]))
    story.append(Spacer(1, 0.2*cm))

    story.append(Paragraph("C. Post-Match Evaluation", S["h2"]))
    for line in _post_match_lines(cfg):
        story.append(Paragraph(f"• {line}", S["bullet"]))
    story.append(Spacer(1, 0.4*cm))

    story.append(Paragraph(
        f"Report generated by Real Madrid CF Analytics Dashboard · "
        f"Report ID: {cfg['report_id']} · "
        f"{cfg['generated_at'].strftime('%d %B %Y %H:%M')}",
        ParagraphStyle("footer2", fontName="Helvetica-Oblique", fontSize=8,
                       textColor=_GRAY_RL, alignment=TA_CENTER)
    ))

    doc.build(story)
    return buf.getvalue()


# ═════════════════════════════════════════════════════════════════════════════
# DOCX GENERATION
# ═════════════════════════════════════════════════════════════════════════════

def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour in DOCX."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)


def _docx_add_heading(doc: Document, text: str, level: int,
                       color: RGBColor | None = None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = color


def _docx_add_para(doc: Document, text: str, bold: bool = False,
                    color: RGBColor | None = None, size_pt: int = 10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold  = bold
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    return p


def _docx_add_table(doc: Document, df: pd.DataFrame, max_rows: int = 40):
    if df is None or df.empty:
        return
    df = df.head(max_rows)
    gold_rgb = RGBColor(*_GOLD_RGB)
    navy_rgb = RGBColor(*_NAVY_RGB)
    t = doc.add_table(rows=1 + len(df), cols=len(df.columns))
    t.style = "Table Grid"
    # Header row
    hdr_cells = t.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
        _set_cell_bg(hdr_cells[i], "0B1730")
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(*_WHITE_RGB)
            run.font.size = Pt(8)
    # Data rows
    for ri, (_, row) in enumerate(df.iterrows()):
        row_cells = t.rows[ri + 1].cells
        bg = "F8FAFC" if ri % 2 == 0 else "FFFFFF"
        for ci, col in enumerate(df.columns):
            row_cells[ci].text = str(row[col])
            _set_cell_bg(row_cells[ci], bg)
            for run in row_cells[ci].paragraphs[0].runs:
                run.font.size = Pt(8)
    doc.add_paragraph()  # spacer after table


def _docx_embed_fig(doc: Document, fig, caption: str = "", width_inch: float = 6.0):
    png = _fig_to_png(fig, width=720, height=380)
    if png:
        buf = io.BytesIO(png)
        doc.add_picture(buf, width=Inches(width_inch))
        if caption:
            cap_p = doc.add_paragraph(caption)
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_p.runs[0].font.size  = Pt(8)
            cap_p.runs[0].font.italic = True
        doc.add_paragraph()


def generate_docx(cfg: dict) -> bytes:
    """Build and return a complete DOCX report as bytes."""
    doc  = Document()
    buf  = io.BytesIO()
    gold = RGBColor(*_GOLD_RGB)
    navy = RGBColor(*_NAVY_RGB)
    gray = RGBColor(*_GRAY_RGB)

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)

    # ── COVER ────────────────────────────────────────────────────────────────
    cover_title = doc.add_heading("REAL MADRID CF", 0)
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cover_title.runs:
        run.font.color.rgb = navy
        run.font.bold = True
        run.font.size = Pt(28)

    sub = doc.add_paragraph("Tactical & Player Analytics Report")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size  = Pt(14)
    sub.runs[0].font.bold  = False
    sub.runs[0].font.color.rgb = gray

    doc.add_paragraph()
    meta_pairs = [
        ("Competition",     cfg["competition"]),
        ("Season",          cfg["season"]),
        ("Match Selection", cfg["match_mode"]),
        ("Matches",         ", ".join(cfg["match_labels"]) or "All"),
        ("Athletes",        ", ".join(cfg["players"]) if cfg["players"] else "All Squad"),
        ("Report ID",       cfg["report_id"]),
        ("Generated",       cfg["generated_at"].strftime("%d %B %Y  %H:%M")),
    ]
    t = doc.add_table(rows=len(meta_pairs), cols=2)
    t.style = "Table Grid"
    for ri, (k, v) in enumerate(meta_pairs):
        t.rows[ri].cells[0].text = k
        t.rows[ri].cells[1].text = v
        _set_cell_bg(t.rows[ri].cells[0], "0B1730")
        for run in t.rows[ri].cells[0].paragraphs[0].runs:
            run.font.color.rgb = gold
            run.font.bold = True
            run.font.size = Pt(10)
        for run in t.rows[ri].cells[1].paragraphs[0].runs:
            run.font.size = Pt(10)
    doc.add_page_break()

    # ── SECTION 1: FILTERS ──────────────────────────────────────────────────
    _docx_add_heading(doc, "Section 1 — Filter Summary", 1, navy)
    filter_df = pd.DataFrame(meta_pairs, columns=["Filter", "Value"])
    _docx_add_table(doc, filter_df)

    # ── SECTION 2: EXECUTIVE SUMMARY ────────────────────────────────────────
    _docx_add_heading(doc, "Section 2 — Executive Summary", 1, navy)
    for line in _exec_summary_lines(cfg):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(line)
        run.font.size = Pt(10)
    doc.add_paragraph()

    # ── SECTION 3: TEAM PERFORMANCE ─────────────────────────────────────────
    _docx_add_heading(doc, "Section 3 — Team Performance Analysis", 1, navy)
    recs = cfg["match_records"]
    if recs:
        n   = len(recs)
        gf  = sum(r["meta"].get("rm_score", 0) for r in recs)
        ga  = sum(r["meta"].get("opp_score", 0) for r in recs)
        xgf = sum(r["kpis"].get("xg_for",       0) for r in recs)
        xga = sum(r["kpis"].get("xg_against",   0) for r in recs)
        sht = sum(r["kpis"].get("shots_total",   0) for r in recs)
        poss_v = [r["kpis"].get("possession",   0) for r in recs if r["kpis"].get("possession")]
        pacc_v = [r["kpis"].get("pass_accuracy",0) for r in recs if r["kpis"].get("pass_accuracy")]
        ppda_v = [r["ppda"] for r in recs if isinstance(r.get("ppda"), (int, float))]
        team_df = pd.DataFrame({
            "Metric":    ["Matches","Goals","Conceded","xG For","xG Against",
                          "xG Diff","Shots","Avg Poss%","Avg Pass Acc%","Avg PPDA"],
            "Total/Avg": [n, gf, ga, f"{xgf:.2f}", f"{xga:.2f}", f"{xgf-xga:+.2f}",
                          sht,
                          f"{sum(poss_v)/len(poss_v):.1f}%" if poss_v else "N/A",
                          f"{sum(pacc_v)/len(pacc_v):.1f}%" if pacc_v else "N/A",
                          f"{sum(ppda_v)/len(ppda_v):.2f}" if ppda_v else "N/A"],
            "Per Match": [1, f"{gf/n:.2f}", f"{ga/n:.2f}",
                          f"{xgf/n:.2f}", f"{xga/n:.2f}", f"{(xgf-xga)/n:+.2f}",
                          f"{sht/n:.1f}", "—", "—", "—"],
        })
        _docx_add_table(doc, team_df)

        _docx_add_heading(doc, "Match-by-Match Results", 2, gold)
        match_rows = []
        for r in recs:
            meta, kpis = r["meta"], r["kpis"]
            res = ("W" if meta.get("rm_score",0) > meta.get("opp_score",0) else
                   "D" if meta.get("rm_score",0) == meta.get("opp_score",0) else "L")
            match_rows.append({
                "Match":  f"{meta.get('home_team','?')} vs {meta.get('away_team','?')}",
                "Score":  meta.get("score_str","?"),
                "Result": res,
                "Poss%":  f"{kpis.get('possession',0):.0f}%",
                "PA%":    f"{kpis.get('pass_accuracy',0):.0f}%",
                "xGF":    f"{kpis.get('xg_for',0):.2f}",
                "xGA":    f"{kpis.get('xg_against',0):.2f}",
                "PPDA":   str(r.get("ppda","?")),
            })
        _docx_add_table(doc, pd.DataFrame(match_rows))
    else:
        _docx_add_para(doc, "No team performance data for selected filters.")
    doc.add_paragraph()

    # ── SECTION 4: PLAYER PERFORMANCE ───────────────────────────────────────
    _docx_add_heading(doc, "Section 4 — Player Performance Analysis", 1, navy)
    if cfg["players"] and cfg["player_dfs"]:
        for pname in cfg["players"]:
            _docx_add_heading(doc, pname, 2, gold)
            pdf = cfg["player_dfs"].get(pname)
            if pdf is None or pdf.empty:
                _docx_add_para(doc, "No per-match data available.")
                continue
            nm  = len(pdf)
            agg = pdf[["passes","shots","goals","tackles","xg"]].sum()
            pg  = (agg / max(nm, 1)).round(3)
            p_df = pd.DataFrame({
                "Metric":    ["Matches","Passes","Shots","Goals","Tackles","xG"],
                "Total":     [nm, int(agg.get("passes",0)), int(agg.get("shots",0)),
                              int(agg.get("goals",0)), int(agg.get("tackles",0)),
                              f"{float(agg.get('xg',0)):.2f}"],
                "Per Match": ["—", f"{pg.get('passes',0):.1f}", f"{pg.get('shots',0):.2f}",
                              f"{pg.get('goals',0):.2f}", f"{pg.get('tackles',0):.2f}",
                              f"{float(pg.get('xg',0)):.3f}"],
            })
            _docx_add_table(doc, p_df)
            # Embed radar
            radar_fig = cfg["figures"].get(f"radar_{pname}") or cfg["figures"].get("radar_comparison")
            _docx_embed_fig(doc, radar_fig,
                            f"{pname} — Radar Profile (% of squad maximum per metric)")
    elif not cfg["squad_df"].empty:
        _docx_add_heading(doc, "Squad Performance Summary", 2, gold)
        sq = cfg["squad_df"]
        show_cols = [c for c in ["player_name","matches_played","goals","assists",
                                  "passes","shots","tackles","xg"] if c in sq.columns]
        _docx_add_table(doc, sq[show_cols].sort_values("goals", ascending=False).head(15))
    else:
        _docx_add_para(doc, "No player data available for selected filters.")
    doc.add_paragraph()

    # ── SECTION 5: TACTICAL ANALYSIS ────────────────────────────────────────
    _docx_add_heading(doc, "Section 5 — Tactical Analysis", 1, navy)
    tac_figs = {k: v for k, v in cfg["figures"].items()
                if k in ("press_map","press_class","position_map","tactical_bars")}
    if tac_figs:
        for fig_key, fig in tac_figs.items():
            label = fig_key.replace("_", " ").title()
            _docx_add_heading(doc, label, 2, gold)
            _docx_embed_fig(doc, fig, f"{label} — derived from Opta event data")
    else:
        _docx_add_para(doc, "Tactical visualisations available in the dashboard Tactical Phases section.")
    doc.add_paragraph()

    # ── SECTION 6: MATCH ANALYSIS ────────────────────────────────────────────
    _docx_add_heading(doc, "Section 6 — Match Analysis", 1, navy)
    for line in _match_analysis_lines(cfg):
        if line.startswith("  "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line.strip()).font.size = Pt(9)
        else:
            _docx_add_para(doc, line, size_pt=10)
    for fig_key in ("shot_map","xg_chart","pass_map","pass_network"):
        fig = cfg["figures"].get(fig_key)
        if fig:
            _docx_add_heading(doc, fig_key.replace("_"," ").title(), 2, gold)
            _docx_embed_fig(doc, fig, f"{fig_key.replace('_',' ').title()} — selected matches")
    doc.add_paragraph()

    # ── SECTION 7: BENCHMARKING ──────────────────────────────────────────────
    _docx_add_heading(doc, "Section 7 — Benchmarking Analysis", 1, navy)
    bm_fig = cfg["figures"].get("benchmarking_goals") or cfg["figures"].get("benchmarking_xg")
    _docx_embed_fig(doc, bm_fig, "Season benchmarking — aggregated competition data")
    if not bm_fig:
        _docx_add_para(doc, "Benchmarking charts available in the dashboard Benchmarking section.")
    doc.add_paragraph()

    # ── SECTION 8: APPENDIX ──────────────────────────────────────────────────
    _docx_add_heading(doc, "Section 8 — Visualisation Appendix", 1, navy)
    skip = {"shot_map","xg_chart","pass_map","pass_network",
            "press_map","press_class","position_map","tactical_bars",
            "benchmarking_goals","benchmarking_xg"}
    included = 0
    for fig_key, fig in cfg["figures"].items():
        if fig_key in skip or fig is None:
            continue
        label = fig_key.replace("_", " ").title()
        _docx_add_heading(doc, label, 2, gold)
        _docx_embed_fig(doc, fig, label)
        included += 1
    if included == 0:
        _docx_add_para(doc, "No additional visualisations captured for this report.")
    doc.add_paragraph()

    # ── CONCLUSION ───────────────────────────────────────────────────────────
    _docx_add_heading(doc, "Conclusion", 1, navy)

    _docx_add_heading(doc, "A. Pre-Match Analysis", 2, gold)
    for line in _pre_match_lines(cfg):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(line).font.size = Pt(10)

    _docx_add_heading(doc, "B. Match Analysis", 2, gold)
    for line in _match_analysis_lines(cfg):
        if line.startswith("  "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line.strip()).font.size = Pt(9)
        else:
            _docx_add_para(doc, line, size_pt=10)

    _docx_add_heading(doc, "C. Post-Match Evaluation", 2, gold)
    for line in _post_match_lines(cfg):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(line).font.size = Pt(10)

    doc.add_paragraph()
    footer_p = doc.add_paragraph(
        f"Report ID: {cfg['report_id']}  ·  "
        f"Real Madrid CF Analytics Dashboard  ·  "
        f"Generated {cfg['generated_at'].strftime('%d %B %Y %H:%M')}"
    )
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.runs[0].font.size   = Pt(8)
    footer_p.runs[0].font.italic = True
    footer_p.runs[0].font.color.rgb = gray

    doc.save(buf)
    return buf.getvalue()
