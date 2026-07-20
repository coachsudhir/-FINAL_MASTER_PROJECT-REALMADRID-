"""
generate_thesis_docx.py
Builds the submission-ready Master's thesis as a formatted Word document:
  Real Madrid Tactical Dashboard — Sports Analytics Master's Final Project

Output: Real_Madrid_Tactical_Dashboard_Thesis.docx
"""
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand palette ────────────────────────────────────────────────────────────
NAVY = RGBColor(0x0B, 0x17, 0x30)
GOLD = RGBColor(0xA8, 0x88, 0x2E)
GREY = RGBColor(0x47, 0x55, 0x69)
DARK = RGBColor(0x1F, 0x2A, 0x3D)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

OUT = Path(__file__).parent / "Real_Madrid_Tactical_Dashboard_Thesis.docx"

doc = Document()

# ── Base style ───────────────────────────────────────────────────────────────
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = DARK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for lvl, (sz, col) in {1: (16, NAVY), 2: (13, GOLD), 3: (11.5, NAVY)}.items():
    st = doc.styles[f"Heading {lvl}"]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.bold = True
    st.font.color.rgb = col
    st.paragraph_format.space_before = Pt(12 if lvl == 1 else 8)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.keep_with_next = True


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def set_cell(cell, text, *, bold=False, color=None, size=9.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Calibri"
    if color is not None:
        run.font.color.rgb = color


def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        shade(c, "0B1730")
        set_cell(c, h, bold=True, color=WHITE, size=9.5)
    for r in rows:
        cells = t.add_row().cells
        for i, val in enumerate(r):
            set_cell(cells[i], val, size=9)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def body(text, *, italic=False, size=11, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    run.font.italic = italic
    run.font.size = Pt(size)
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)
    return p


def mono(lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(lines)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = NAVY


def page_break():
    doc.add_page_break()


# ── Figures (rendered live from real data by generate_thesis_figures.py) ──────
FIG_DIR = Path(__file__).parent / "thesis_figures"
_fig_n = [0]


def add_figure(filename, caption, width=6.2):
    """Embed a figure with a numbered caption; skip gracefully if missing."""
    path = FIG_DIR / filename
    if not path.exists():
        print(f"  ! figure missing, skipped: {filename} (run generate_thesis_figures.py)")
        return
    _fig_n[0] += 1
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(f"Figure {_fig_n[0]}. {caption}")
    r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = GREY


# ═══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("REAL MADRID TACTICAL DASHBOARD")
r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = NAVY
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Transforming Event-Level Football Data into Interactive Tactical Intelligence")
r.font.size = Pt(14); r.font.italic = True; r.font.color.rgb = GOLD

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Master's Thesis — Final Project")
r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = GREY
for _ in range(2):
    doc.add_paragraph()

meta = [
    ("Student", "Sudhir Dahiya"),
    ("Degree", "Master's in Sports Analytics (2025–2026)"),
    ("Institution", "Escuela Universitaria Real Madrid — Universidad Europea"),
    ("Case Club", "Real Madrid CF"),
    ("Academic Year", "2025–2026"),
    ("Data Provider", "Opta Stats Perform (event-level)"),
    ("Technology Stack", "Python · Pandas · NumPy · Plotly · Streamlit (Dash analytics core)"),
]
tbl = doc.add_table(rows=0, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for k, v in meta:
    cells = tbl.add_row().cells
    set_cell(cells[0], k, bold=True, color=NAVY, size=11, align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell(cells[1], v, size=11)
    cells[0].width = Inches(2.2); cells[1].width = Inches(4.0)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
r = p.add_run(f"Generated {date.today():%d %B %Y}")
r.font.size = Pt(9); r.font.color.rgb = GREY
page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 2 · ABSTRACT
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("Abstract", 1)
body(
    "Elite football decision-making increasingly depends on the ability to convert vast, "
    "granular event data into tactical understanding faster than an opponent can. This project "
    "addresses a persistent gap in applied football analytics: while event data is abundant, "
    "integrated tools that translate it into phase-specific tactical insight for coaching staff "
    "remain scarce. The work presents the Real Madrid Tactical Dashboard, an interactive "
    "analytical platform that ingests Opta Stats Perform event-level JSON for 108 Real Madrid CF "
    "matches across three competitions (LaLiga, UEFA Champions League, Copa del Rey) and two "
    "seasons (2024–25, 2025–26), and transforms them into match-, season-, and player-level "
    "tactical intelligence.")
body(
    "The methodology rests on a reproducible Python pipeline that normalises raw event streams "
    "and computes a transparent suite of KPIs — including a calibrated positional Expected Goals "
    "(xG) model, PPDA (pressing intensity), field tilt, and a four-phase tactical model "
    "(offensive moment, defensive transition, defensive moment, offensive transition) — rendered "
    "through an interactive Streamlit interface backed by Plotly visualisations. All metrics are "
    "derived strictly from observed event data, with no synthetic values.")
body(
    "The resulting system enables analysts to interrogate how Real Madrid plays rather than merely "
    "what the scoreline was: quantifying transition threat, pressing structure, attacking "
    "organisation, and performance against expectation. The dashboard demonstrates measurable value "
    "for pre-match preparation, post-match review, and opposition scouting, and establishes an "
    "extensible foundation for predictive and player-level analytics. The contribution is both "
    "practical — a club-grade tactical tool — and academic: a documented, formula-transparent "
    "framework for phase-based football analysis.")
page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 3 · INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("1 · Introduction", 1)
doc.add_heading("1.1 Football Analytics Background", 2)
body(
    "Football has transitioned from a results-recording discipline to a data-rich performance "
    "science. Modern providers such as Opta Stats Perform capture every on-ball event — passes, "
    "shots, duels, recoveries, set-piece restarts — with spatial coordinates and qualifier "
    "metadata, producing thousands of structured events per match. This density enables analytical "
    "questions previously reserved for subjective observation: where a team regains the ball, how "
    "quickly it converts regains into shots, and whether its shot volume is sustainable relative to "
    "chance quality.")
doc.add_heading("1.2 The Importance of Tactical Analysis", 2)
body(
    "Results are noisy; tactics are signal. A single deflected goal can decide a match, but the "
    "underlying tactical behaviours — pressing triggers, build-up structure, transition speed, "
    "territorial dominance — are stable, repeatable, and coachable. Quantifying these behaviours "
    "allows staff to validate that the team is executing its intended game model, diagnose "
    "deviations, and prepare specifically against an opponent's structural tendencies. Expected "
    "Goals and pressing metrics such as PPDA have become industry standards precisely because they "
    "correlate with sustained performance better than raw outcomes.")
doc.add_heading("1.3 Why Real Madrid CF as a Case Study", 2)
body(
    "Real Madrid CF is an analytically compelling subject. The club competes simultaneously across "
    "LaLiga, the Champions League, and the Copa del Rey, generating a multi-competition dataset that "
    "exposes tactical adaptation across opponent quality tiers. Its identity blends possession "
    "dominance with devastating transition play, making it an ideal vehicle to demonstrate "
    "phase-based analysis, where the same team must be characterised differently in settled "
    "possession versus moments of transition. As the host institution's namesake club, it also "
    "grounds the project in a realistic elite-environment workflow.")

# ═══════════════════════════════════════════════════════════════════════════
# 4 · PROBLEM STATEMENT
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("2 · Problem Statement", 1)
body("Traditional football analysis, and even much contemporary data work, suffers from three "
     "structural limitations that this project targets:")
bullet("Outcome bias over behaviour. Conventional reporting emphasises goals, possession "
       "percentage, and results, which describe what happened but not how or why. A 2–0 win "
       "conceals whether it was earned through sustained control or fortunate finishing of "
       "low-quality chances.")
bullet("Fragmentation of tactical tools. Analysts typically stitch together spreadsheets, isolated "
       "notebooks, and static slide decks. There is rarely a single, interactive environment that "
       "links season context, match detail, tactical phases, and player contribution within one "
       "consistent, reproducible data model.")
bullet("The data-to-decision gap. Raw event feeds are inaccessible to most coaching staff. "
       "Translating a JSON event stream into an actionable tactical statement requires an "
       "engineering and modelling layer that most clubs lack in a transparent, auditable form.")
body("The problem this thesis addresses: the absence of an integrated, transparent, event-driven "
     "tactical dashboard that quantifies Real Madrid's playing identity in phase-specific terms and "
     "presents it in a form usable by analysts and coaches.", italic=True)

# ═══════════════════════════════════════════════════════════════════════════
# 5 · OBJECTIVES
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("3 · Objectives", 1)
body("The project pursues four explicit, measurable objectives:")
for txt in [
    "Quantify tactical behaviour. Convert raw event data into defensible numerical descriptors of "
    "Real Madrid's pressing, possession, transition, and attacking behaviour, using "
    "formula-transparent KPIs.",
    "Implement phase-based match analysis. Segment each match into the four canonical tactical "
    "phases and score the team's effectiveness within each.",
    "Model performance versus expectation. Build a calibrated positional xG model to compare actual "
    "output against chance quality, distinguishing sustainable performance from variance.",
    "Deliver an interactive dashboard. Engineer a reproducible Python pipeline and an interactive "
    "interface that lets users explore season, match, tactical, and player views without writing code.",
]:
    bullet(txt)

# ═══════════════════════════════════════════════════════════════════════════
# 6 · DATASET & SCOPE
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("4 · Dataset & Scope", 1)
doc.add_heading("4.1 Source and Structure", 2)
body("The dataset comprises Opta Stats Perform event-level JSON match files for every Real Madrid CF "
     "fixture across the covered competitions and seasons. Each file contains a complete event stream "
     "(typed events with x/y pitch coordinates, outcomes, player and team identifiers, and qualifier "
     "metadata) plus match metadata (teams, score, venue, competition, stage).")
doc.add_heading("4.2 Scope", 2)
add_table(
    ["Dimension", "Coverage"],
    [["Competitions", "LaLiga · UEFA Champions League · Copa del Rey"],
     ["Seasons", "2024–25, 2025–26"],
     ["Total clean match files", "108 (verified, 0 corrupt)"],
     ["2025–26 matches", "50 (LaLiga 36 · UCL 12 · Copa 2)"],
     ["Event types used", "Passes, shots, tackles, interceptions, recoveries, dribbles, fouls, corners, cards, subs"],
     ["Spatial resolution", "Opta 0–100 normalised pitch coordinates"],
     ["Granularity", "Single-event (on-ball)"]],
    widths=[2.2, 4.3])
doc.add_heading("4.3 Case-Study Lens & Boundaries", 2)
body("While the platform indexes the full 108-match corpus, the tactical case study foregrounds the "
     "LaLiga 2025–26 campaign as the primary competitive context, with a rolling recent-match window "
     "used to illustrate match-level drill-downs. The dataset is event (on-ball) data only; it contains "
     "no optical/GPS tracking, so possession is represented as a pass-share proxy rather than a "
     "stopwatch duration.")
page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 7 · METHODOLOGY
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("5 · Methodology", 1)
doc.add_heading("5.1 Event-Based Data Processing", 2)
body("Each match JSON is parsed into a normalised, typed event table. The pipeline maps Opta typeIds "
     "to semantic flags (is_pass, is_shot, is_tackle, is_recovery, etc.), extracts pass end-coordinates "
     "and set-piece qualifiers, removes bookkeeping events (period starts, formation markers), and "
     "attaches a per-shot xG value. Results are LRU-cached for responsiveness.")
doc.add_heading("5.2 KPI Computation", 2)
add_table(
    ["KPI", "Definition"],
    [["Possession %", "RM passes ÷ (RM + opponent passes) × 100"],
     ["Pass accuracy %", "mean(outcome) × 100 over RM passes"],
     ["xG (positional)", "logistic of shot distance & angle; penalties fixed at 0.76"],
     ["PPDA", "opponent passes (x ≥ 40) ÷ RM defensive actions (tackle/interception/recovery, x ≥ 40)"],
     ["Field tilt", "RM final-third touches (x ≥ 67) ÷ all final-third touches × 100"]],
    widths=[1.6, 4.9])
body("Positional xG model:", after=2)
mono("dx = (100 − x)·1.05 ;  dy = (y − 50)·0.68\n"
     "dist = √(dx² + dy²)\n"
     "angle = | atan2(dy − 3.66, dx) − atan2(dy + 3.66, dx) |\n"
     "logit(xG) = −3.785 − 0.0337·dist + 3.64·angle\n"
     "xG = 1 / (1 + e^−logit),  clipped to [0.01, 0.99]")
doc.add_heading("5.3 Tactical Phase Segmentation", 2)
body("Four phases are scored on a normalised 0–100 scale, blending weighted components:")
mono("A · Offensive Moment      = pass_acc·0.45 + (Σxg·10)·0.35 + (shots·2.5)·0.20\n"
     "B · Defensive Transition  = (100/(1+PPDA))·0.55 + ((recov+inter)·1.8)·0.45\n"
     "C · Defensive Moment      = (100/(1+PPDA))·0.40 + max(0, 40 − opp_shots·2)·0.60\n"
     "D · Offensive Transition  = trans_rate·0.65 + (Σxg·7)·0.35")
body("where trans_rate = the percentage of ball regains (recovery/interception/tackle) that produce a "
     "Real Madrid shot within 15 seconds — a direct, event-timed measure of transition threat.")
doc.add_heading("5.4 Python Analysis Pipeline", 2)
body("The pipeline is deterministic and reproducible: JSON → parse_events → calc_match_kpis → "
     "calc_season_kpis / phase_scores / player_stats → visual + report layer. All numerical logic is "
     "isolated in utility modules so it can be validated independently of the interface.")

# ═══════════════════════════════════════════════════════════════════════════
# 8 · SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("6 · System Architecture", 1)
mono("OPTA STATS PERFORM (event JSON)\n"
     "        ↓\n"
     "data_loader.py  → parse & normalise events (Pandas / NumPy)\n"
     "        ↓\n"
     "KPI & MODEL LAYER → xG · PPDA · field tilt · phase scoring · player stats\n"
     "        ↓\n"
     "VISUALISATION LAYER → Plotly (shot maps, heatmaps, radars, trends)\n"
     "        ↓\n"
     "PRESENTATION LAYER → Streamlit UI + Dash analytics core\n"
     "        ↓\n"
     "EXPORT LAYER → automated PDF / DOCX reporting")
add_table(
    ["Layer", "Technology", "Role"],
    [["Data processing", "Python, Pandas, NumPy", "Parsing, normalisation, aggregation"],
     ["Modelling", "Python (math/SciPy)", "xG, PPDA, phase scoring"],
     ["Visualisation", "Plotly", "Interactive pitch maps, charts, radars"],
     ["Interface", "Streamlit (front-end), Dash (heritage)", "Interactive multi-page exploration"],
     ["Reporting", "ReportLab, python-docx", "Automated match/season report export"],
     ["Deployment", "Render (CI/CD from GitHub)", "Public, auto-deployed web service"]],
    widths=[1.4, 2.3, 2.8])
body("Data flow: a user selects competition → season → match; the relevant JSON is loaded (cached), "
     "parsed, scored, and rendered. Season views aggregate across all matches in a competition-season; "
     "player and benchmarking views re-slice the same normalised events.")
page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 9 · DASHBOARD MODULES
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("7 · Dashboard Modules (Core Section)", 1)
body("Each module is documented by description, metrics, visualisation, tactical insight, and current "
     "status. The requested module taxonomy is mapped onto the implemented multi-page architecture.")

modules = [
    ("7.1 Pre-Match Analysis",
     "Establishes season baselines and opponent context before a fixture.",
     "Goals scored/conceded, xG vs xGA, win rate, possession & pass-accuracy trends, opponent record.",
     "KPI cards, season trend lines, results table, opponent profile.",
     "Frames whether recent results are sustainable (xG-aligned) and characterises the opponent.",
     "Implemented (Home/Season Overview + Opponent Analysis + Benchmarking)."),
    ("7.2 Attacking Transitions",
     "Quantifies threat generated immediately after winning the ball.",
     "Transition rate (regain→shot ≤15 s), Phase-D score, xG from transition sequences.",
     "Phase-D scoring; fast-break shot locations.",
     "Measures Real Madrid's hallmark vertical transition threat in objective terms.",
     "Implemented — Transition Metrics, Extended Transition Analysis, A/B/C/D sub-phases, "
     "configurable 5–15 s regain→shot window."),
    ("7.3 Organized Attack",
     "Analyses structured, settled-possession attacking.",
     "Shot count/quality, xG, pass volume & accuracy, field tilt.",
     "Shot maps, pass networks, attacking-third activity.",
     "Distinguishes chance creation through organised build-up from transition output.",
     "Implemented — Shot Map, Shot Zone Map, Pass Map, Build-Up Network, Zone 14, Progressive Passes, "
     "Chance Creation Heatmap."),
    ("7.4 Defensive Strategy",
     "Characterises pressing and defensive solidity.",
     "PPDA, opponent shots/xG conceded, tackles, interceptions, recoveries.",
     "PPDA trend, pressing maps, defensive-action distribution.",
     "Reveals pressing intensity and where the team concedes territory/chances.",
     "Implemented — PPDA Trend, Pressing Actions Map, Defensive Actions Map (shots-conceded heatmap "
     "is future work)."),
    ("7.5 Set Piece Analysis",
     "Isolates dead-ball contribution.",
     "Set-piece-originated shots/goals (restart qualifiers 5/72).",
     "Goals-from-set-pieces bar chart.",
     "Assesses dependence on and efficiency of set plays.",
     "Implemented — Set Piece Efficiency panel; set-piece tendencies (corners, shots ≤20 s) in "
     "Opponent Scout."),
    ("7.6 Possession Recovery",
     "Maps where and how the team regains possession.",
     "Ball recoveries (typeId 49), interceptions, recovery zones.",
     "Recovery heatmap, zone-based aggregation.",
     "Indicates whether regains feed transitions (high) or rebuilds (deep).",
     "Implemented — Ball Recoveries by Zone (zonal heatmap)."),
    ("7.7 Final Third Entry Analysis",
     "Quantifies territorial penetration.",
     "Final-third touches (x ≥ 67), field tilt %.",
     "Final-third entry heatmap.",
     "Measures sustained territorial dominance vs sporadic penetration.",
     "Implemented — Final Third Entry Analysis panel; Field Tilt card &amp; season trend (x ≥ 67)."),
    ("7.8 Post-Match Analysis",
     "Reviews performance against expectation.",
     "Goals vs xG, xGA vs goals conceded, defensive actions, recoveries.",
     "xG accumulation curve, KPI deltas, defensive summaries.",
     "Separates earned performance from finishing/variance.",
     "Implemented (post-match KPIs + xG accumulation)."),
    ("7.9 Match Information",
     "Contextual match header.",
     "Score, opponent, venue, competition, result.",
     "Summary header / metadata block.",
     "Anchors every analysis in its competitive context.",
     "Implemented (extract_match_meta)."),
]
for title, desc, metrics, vis, insight, status in modules:
    doc.add_heading(title, 3)
    for lbl, val in [("Description", desc), ("Metrics", metrics),
                     ("Visualisation", vis), ("Tactical insight", insight),
                     ("Status", status)]:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{lbl}: "); r.font.bold = True; r.font.color.rgb = NAVY; r.font.size = Pt(10)
        r2 = p.add_run(val); r2.font.size = Pt(10)
page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 10 · TACTICAL INSIGHTS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("8 · Tactical Insights — Real Madrid Case Study", 1)
body("Figures below are the dashboard's computed outputs for the LaLiga 2025–26 campaign, presented to "
     "illustrate the platform's analytical reach. xG figures derive from the project's positional model, "
     "not the provider's proprietary xG.", italic=True, size=10)
doc.add_heading("8.1 Transition-Heavy Attacking Identity", 2)
body("Real Madrid's offensive signature is the speed of conversion from regain to shot. The Phase-D "
     "(offensive transition) model — anchored on the proportion of regains producing a shot within 15 "
     "seconds — consistently scores the side highly, confirming quantitatively what is observed "
     "qualitatively: the team is most dangerous in the moments immediately after winning the ball.")
add_figure("fig_phase_radar.png",
           "Four-phase tactical profile (A/B/C/D, normalised 0–100) for Real Madrid vs Barcelona "
           "(LaLiga 2025–26), computed live from Opta event data via the dashboard's phase-scoring model. "
           "The pronounced offensive-moment axis reflects the side's settled-attack productivity in this fixture.")
doc.add_heading("8.2 Possession With Penetration", 2)
body("Across LaLiga 2025–26 the dashboard reports approximately 58.8% possession and 87.7% pass "
     "accuracy, yet the value lies in pairing these with field tilt and final-third activity: control is "
     "converted into territory rather than sterile circulation. Shot generation is distributed between "
     "settled build-up and transition, indicating a dual-threat profile.")
doc.add_heading("8.3 Efficiency vs Expectation", 2)
body("With a season xG-for of about 66.0 against 72 goals scored, the team marginally outperforms chance "
     "quality (a goals/xG ratio just above 1.0), consistent with elite finishing — but the dashboard "
     "correctly flags this as a watch-point: sustained over-performance regresses, so the xGA side "
     "(about 38.6 vs 33 conceded) is the more reassuring structural indicator.")
add_figure("fig_shot_map.png",
           "Real Madrid shot map for the showcase fixture, generated from the dashboard's get_shot_data "
           "pipeline. Marker size is proportional to positional-model xG; goals are highlighted. The "
           "concentration of high-xG attempts inside the box illustrates the side's central penetration.")
doc.add_heading("8.4 Defensive Structure", 2)
body("PPDA values (season average about 9.5, with elite single-match presses near 3.3) describe a side "
     "that presses with controlled intensity rather than relentlessly — pressing selectively, then "
     "defending with a compact mid-block. The defensive-moment phase score, driven by low opponent shot "
     "volume, confirms structural solidity outside high-press windows.")
add_figure("fig_ppda_trend.png",
           "Pressing intensity (PPDA) across all 36 LaLiga 2025–26 matchdays, computed per match by the "
           "dashboard's calc_ppda function (lower = more intense press). The dashed line marks the season "
           "average; the variance illustrates the side's selective, game-state-dependent pressing.")
doc.add_heading("8.5 Strengths, Weaknesses & Tactical Identity", 2)
body("Strengths:", after=2)
for t in ["Elite, time-bounded transition threat (objectively quantified).",
          "Possession dominance that translates into territorial control.",
          "Strong chance suppression (low xGA) underpinning results."]:
    bullet(t)
body("Weaknesses / watch-points:", after=2)
for t in ["Mild over-performance versus xG-for (variance risk).",
          "Pressing intensity is selective; lapses in mid-block compactness appear in worst-case PPDA matches.",
          "Reduced creativity against deep, compact defensive blocks."]:
    bullet(t)
body("Tactical identity conclusion: Real Madrid emerges, in data, as a proactive possession side with a "
     "world-class transition engine and disciplined defensive economy — winning through chance "
     "suppression and ruthless conversion of regains, rather than through pressing volume alone.")

# ═══════════════════════════════════════════════════════════════════════════
# 11 · IMPLEMENTATION STATUS TABLE
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("9 · Implementation Status Table", 1)
add_table(
    ["Feature / Module", "Current Status", "Improvements Needed"],
    [["Match information & metadata", "Implemented", "Referee, weather, formation parsing"],
     ["Season KPIs & trends (pre-match)", "Implemented", "Rolling-form windows, opponent-adjusted baselines"],
     ["xG vs xGA / performance vs expectation", "Implemented", "Validate model vs provider xG"],
     ["Shot maps (organised attack)", "Implemented", "xG-weighted markers, body-part split"],
     ["Pass networks", "Implemented", "Edge-weighted lanes, phase filtering"],
     ["PPDA & pressing (defensive strategy)", "Implemented", "Pressing-trigger detection, by-zone PPDA"],
     ["Tactical phase scoring (A/B/C/D)", "Implemented", "Model validation; weighting sensitivity"],
     ["Post-match review (xG curve, def. actions)", "Implemented", "Automated narrative summary"],
     ["Player analysis (squad table, radars)", "Implemented", "Per-90 normalisation, role radars"],
     ["Benchmarking (cross-competition)", "Implemented", "League-percentile context"],
     ["Opponent analysis / scouting", "Implemented", "Opponent tendency modelling"],
     ["Attacking transitions (Transition Metrics)", "Implemented", "Time-to-first-shot distribution"],
     ["Possession recovery (Recoveries by Zone)", "Implemented", "High/mid/low zonal comparison"],
     ["Final-third entries (+ Field Tilt)", "Implemented", "xG-from-entries, shot-probability maps"],
     ["Set-piece analysis (Set Piece Efficiency)", "Implemented", "Dedicated set-piece shot maps"],
     ["Shots-conceded heatmap", "Planned", "Defensive shot-concession map (count shown today)"],
     ["Block-height visualisation", "Planned", "Defensive line-height metric & viz"],
     ["Automated PDF/DOCX reporting", "Implemented", "Templated per-opponent scouting export"],
     ["Cloud deployment (Render CI/CD)", "Implemented", "Caching/CDN for large dataset"]],
    widths=[2.7, 1.3, 2.5])
page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 12 · LIMITATIONS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("10 · Limitations", 1)
for t in [
    "Event data only — no tracking. The pipeline observes on-ball events but not continuous off-ball "
    "positioning, so defensive shape and pressing distances can only be approximated.",
    "Possession is a proxy. Without ball-possession durations, possession% is computed from pass-volume "
    "share — directionally accurate but not a stopwatch figure.",
    "xG is a transparent positional model, not a provider model. It uses shot distance and angle only, "
    "omitting defensive pressure, assist type, and goalkeeper positioning, by design for auditability.",
    "Limited predictive modelling. The system is descriptive and diagnostic; it does not yet forecast "
    "outcomes or simulate tactical scenarios.",
    "No real-time ingestion. Analysis is post-hoc on stored files; there is no live in-match feed.",
    "Phase scores are weighted composites. They rank relative phase strength on a 0–100 scale and are "
    "sensitive to weighting choices; they are interpretive indices, not physical units.",
]:
    bullet(t)

# ═══════════════════════════════════════════════════════════════════════════
# 13 · FUTURE SCOPE
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("11 · Future Scope", 1)
body("Clearly separated from current capability, the roadmap prioritises:")
for t in [
    "Player-level analysis (deepened): per-90 normalised contributions, role-specific radars, on/off "
    "impact, and individual transition involvement.",
    "Multi-season dataset expansion: extend beyond two seasons for longitudinal trend and "
    "tactical-evolution tracking.",
    "Opponent modelling: automated extraction of opponent pressing triggers, build-up patterns, and "
    "set-piece tendencies for on-demand opposition reports.",
    "Predictive analytics: xG/xGA forecasting, expected-points projection, and transition-threat "
    "prediction conditioned on game state.",
    "AI-driven tactical clustering: unsupervised clustering of match phases and possession sequences to "
    "discover recurring patterns and detect anomalies automatically.",
    "Tracking-data integration: fusing optical/GPS tracking for off-ball positioning, pressing "
    "distances, defensive compactness, and pitch control.",
    "Video integration: linking each event/metric to synchronised clip timestamps.",
    "Real-time ingestion: live event streaming for in-match decision support.",
]:
    bullet(t)

# ═══════════════════════════════════════════════════════════════════════════
# 14 · TOOLS & TECHNOLOGIES
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("12 · Tools & Technologies", 1)
add_table(
    ["Category", "Technology", "Purpose"],
    [["Language", "Python 3.11", "Core implementation"],
     ["Data", "Pandas, NumPy", "Event parsing, vectorised aggregation"],
     ["Modelling", "Python math / SciPy", "xG, PPDA, phase scoring"],
     ["Visualisation", "Plotly", "Interactive pitch maps, charts, radars"],
     ["Interface", "Streamlit (Dash analytics core)", "Interactive multi-page dashboard"],
     ["Reporting", "ReportLab, python-docx", "Automated PDF/DOCX export"],
     ["Data format", "Opta Stats Perform event JSON", "Source event data"],
     ["Version control / CI-CD", "Git, GitHub, Render", "Reproducibility and deployment"]],
    widths=[1.8, 2.4, 2.3])

# ═══════════════════════════════════════════════════════════════════════════
# 15 · CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("13 · Conclusion", 1)
body("This project delivers a working, club-grade Real Madrid Tactical Dashboard that closes the "
     "data-to-decision gap for event-level football analysis. By normalising 108 Opta match files into a "
     "reproducible pipeline and computing a transparent, formula-documented suite of KPIs — positional "
     "xG, PPDA, field tilt, and a four-phase tactical model — it characterises how Real Madrid plays, "
     "not merely the result.")
p = doc.add_paragraph(); r = p.add_run("Tactical relevance: "); r.font.bold = True; r.font.color.rgb = NAVY
p.add_run("the platform objectively confirms and quantifies Real Madrid's identity — a possession-dominant "
          "side with an elite, time-bounded transition engine and disciplined chance suppression — and "
          "surfaces actionable watch-points (xG over-performance, selective pressing).")
p = doc.add_paragraph(); r = p.add_run("Value for professional football staff: "); r.font.bold = True; r.font.color.rgb = NAVY
p.add_run("the integrated pre-match, post-match, tactical-phase, and opposition views, combined with "
          "automated reporting, compress an analyst's preparation workflow into a single interactive "
          "environment aligned with elite-club standards.")
p = doc.add_paragraph(); r = p.add_run("Academic contribution: "); r.font.bold = True; r.font.color.rgb = NAVY
p.add_run("the thesis provides a fully documented, auditable methodology for phase-based football "
          "analysis — every metric is traceable to an explicit formula and to observed event data, with "
          "no synthetic values — making the work reproducible and extensible. The delineated roadmap "
          "positions the dashboard as a foundation for continued research and professional deployment "
          "rather than a finished endpoint.")
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Real Madrid CF Tactical & Player Performance Analytics · Data: Opta Stats Perform · "
              "Built with Python, Pandas, NumPy, Plotly and Streamlit\n"
              "Sudhir Dahiya · Master's in Sports Analytics 2025–2026")
r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = GREY


# ═══════════════════════════════════════════════════════════════════════════
# FOOTER WITH PAGE NUMBERS
# ═══════════════════════════════════════════════════════════════════════════
def add_page_number_footer(document):
    footer = document.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Real Madrid Tactical Dashboard — Master's Thesis  ·  Page ")
    run.font.size = Pt(8); run.font.color.rgb = GREY
    # PAGE field
    fld1 = OxmlElement("w:fldSimple"); fld1.set(qn("w:instr"), "PAGE")
    run2 = p.add_run(); run2.font.size = Pt(8); run2.font.color.rgb = GREY
    run2._r.addprevious(fld1)


add_page_number_footer(doc)

doc.save(str(OUT))
print(f"Saved → {OUT}  ({OUT.stat().st_size // 1024} KB)")
